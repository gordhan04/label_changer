import os
from docx import Document

def replace_partial_labels(doc_path, output_path, original_text, new_labels):
    doc = Document(doc_path)
    label_index = 0

    for para in doc.paragraphs:
        for run in para.runs:
            if original_text in run.text:
                if label_index < len(new_labels):
                    run.text = run.text.replace(original_text, new_labels[label_index])
                    label_index += 1
                else:
                    run.text = run.text.replace(original_text, "")

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if original_text in run.text:
                            if label_index < len(new_labels):
                                run.text = run.text.replace(original_text, new_labels[label_index])
                                label_index += 1
                            else:
                                run.text = run.text.replace(original_text, "")

    doc.save(output_path)
    print(f"\n✅ Created: {output_path} with {label_index} labels.")
    os.startfile(output_path)

# --- MAIN PROGRAM ---

print("💡 Leave label empty to stop adding more.")

# Paths
while True:
    template_path = input("Enter template file path(or leave blank for Default): ").strip()
    if not template_path:
        template_path = r"D:\Code\LABEL CHANGER\STICKER.docx"
        break
while True:
    output_path = input("Enter output file path(or leave blank for Default): ").strip()
    if not output_path:
        output_path = r"D:\Code\LABEL CHANGER\newsticker.docx"
        break

# template_path = input("Enter template file path: ").strip()
# output_path = input("Enter output file path(or leave blank for Default): ").strip()
original_text = "LABEL"

new_labels = []

while True:
    label = input("\nEnter new label (or leave blank to finish): ").strip()
    if not label:
        break
    try:
        qty = int(input(f"How many times to repeat '{label}'? "))
        new_labels.extend([label] * qty)
    except ValueError:
        print("❌ Please enter a valid number.")

# Confirm and replace
print(f"\n🔁 Total labels to insert: {len(new_labels)}")
replace_partial_labels(template_path, output_path, original_text, new_labels)
