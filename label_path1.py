import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
import os

def replace_partial_labels(doc_path, output_path, original_text, new_labels):
    doc = Document(doc_path)
    label_index = 0

    for para in doc.paragraphs:
        for run in para.runs:
            if original_text in run.text:
                if label_index < len(new_labels):
                    run.text = run.text.replace(original_text, new_labels[label_index])
                    label_index += 1
                else:
                    run.text = run.text.replace(original_text, "")

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if original_text in run.text:
                            if label_index < len(new_labels):
                                run.text = run.text.replace(original_text, new_labels[label_index])
                                label_index += 1
                            else:
                                run.text = run.text.replace(original_text, "")

    doc.save(output_path)
    messagebox.showinfo("Success", f"✅ Created: {output_path} with {label_index} replaced labels.")
    os.startfile(output_path)

def browse_template():
    file_path = filedialog.askopenfilename(filetypes=[("Word Documents", "*.docx")])
    if file_path:
        template_entry.delete(0, tk.END)
        template_entry.insert(0, file_path)

def add_label():
    label = label_entry.get().strip()
    qty = quantity_entry.get().strip()

    if not label:
        messagebox.showerror("Error", "Label cannot be empty.")
        return
    if not qty.isdigit():
        messagebox.showerror("Error", "Quantity must be a number.")
        return

    new_labels.extend([label] * int(qty))
    label_entry.delete(0, tk.END)
    quantity_entry.delete(0, tk.END)
    update_label_list()

def update_label_list():
    label_listbox.delete(0, tk.END)
    for idx, label in enumerate(new_labels, 1):
        label_listbox.insert(tk.END, f"{idx}. {label}")

def generate():
    template_path = template_entry.get().strip()
    if not os.path.isfile(template_path):
        messagebox.showerror("Error", "Template path is invalid.")
        return

    output_path = template_path.replace(".docx", "_output.docx")
    replace_partial_labels(template_path, output_path, "LABEL", new_labels)

# GUI Setup
root = tk.Tk()
root.title("Label Replacer by Govardhan Raj")
root.geometry("600x500")
root.config(bg="white")

tk.Label(root, text="This app is developed by Govardhan Raj", font=("Arial", 14, "bold"), fg="green", bg="white").pack(pady=10)

frame = tk.Frame(root, bg="white")
frame.pack(pady=10)

# Template Selection
tk.Label(frame, text="Select Word Template (.docx):", bg="white").grid(row=0, column=0, sticky="w")
template_entry = tk.Entry(frame, width=50)
template_entry.grid(row=1, column=0, padx=5, pady=5)
tk.Button(frame, text="Browse", command=browse_template).grid(row=1, column=1, padx=5)

# Label and Quantity Entry
tk.Label(frame, text="Enter Label:", bg="white").grid(row=2, column=0, sticky="w", pady=(10, 0))
label_entry = tk.Entry(frame, width=30)
label_entry.grid(row=3, column=0, pady=5, sticky="w")

tk.Label(frame, text="Quantity:", bg="white").grid(row=2, column=1, sticky="w", pady=(10, 0))
quantity_entry = tk.Entry(frame, width=10)
quantity_entry.grid(row=3, column=1, pady=5, sticky="w")

tk.Button(frame, text="Add Label", command=add_label).grid(row=4, column=0, columnspan=2, pady=10)

# Label List
tk.Label(root, text="Added Labels:", bg="white").pack()
label_listbox = tk.Listbox(root, width=50, height=10)
label_listbox.pack(pady=5)

# Generate Button
tk.Button(root, text="Generate Word File", command=generate, bg="green", fg="white", font=("Arial", 12, "bold")).pack(pady=20)

new_labels = []

root.mainloop()
