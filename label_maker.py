from docx import Document
import os
import ttkbootstrap as tb
from ttkbootstrap.constants import *
from tkinter import messagebox

# ---------------- LOGIC (UNCHANGED) ----------------
def replace_partial_labels(doc_path, output_path, original_text, new_labels):
    doc = Document(doc_path)
    label_index = 0

    for para in doc.paragraphs:
        for run in para.runs:
            if original_text in run.text:
                if label_index < len(new_labels):
                    run.text = run.text.replace(original_text, new_labels[label_index])
                    label_index += 1
                else:
                    run.text = run.text.replace(original_text, "")

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if original_text in run.text:
                            if label_index < len(new_labels):
                                run.text = run.text.replace(original_text, new_labels[label_index])
                                label_index += 1
                            else:
                                run.text = run.text.replace(original_text, "")

    doc.save(output_path)
    # messagebox.showinfo("Done", f"✅ Created: {output_path}")
    os.startfile(output_path)


template_path = r"D:\Code\LABEL CHANGER\STICKER.docx"
output_path = r"D:\Code\LABEL CHANGER\newsticker.docx"
original_text = "LABEL"

new_labels = []

def add_label():
    label = label_entry.get().strip()
    qty = qty_entry.get().strip()

    if not label:
        messagebox.showwarning("Input Error", "Label cannot be blank")
        return

    try:
        qty = int(qty)
        if not (1 <= qty <= 48):
            raise ValueError
        new_labels.extend([label] * qty)
        label_entry.delete(0, END)
        qty_entry.delete(0, END)
        total_label.config(text=f"Total Labels: {len(new_labels)}")
    except ValueError:
        messagebox.showerror("Input Error", "Quantity must be between 1 and 48")

def run_replace():
    if not new_labels:
        messagebox.showwarning("No Labels", "Please add labels first")
        return
    replace_partial_labels(template_path, output_path, original_text, new_labels)

def reset_app():
    new_labels.clear()
    label_entry.delete(0, END)
    qty_entry.delete(0, END)
    total_label.config(text="Total Labels: 0")

# ---------------- MODERN UI ----------------
app = tb.Window(themename="superhero")   # try: cosmo, flatly, litera, superhero
app.title("Label Changer Pro")
app.geometry("480x600")
app.resizable(False, False)

# Center window
app.place_window_center()

# Header
header = tb.Label(
    app,
    text="Label Changer",
    font=("Segoe UI", 26, "bold"),
    bootstyle=INFO
)
header.pack(pady=15)

sub = tb.Label(
    app,
    text="Developed by Govardhan Raj",
    font=("Segoe UI", 11),
    bootstyle=SECONDARY
)
sub.pack()

# Card Frame
card = tb.Frame(app, padding=25, bootstyle="secondary")
card.pack(padx=20, pady=25, fill=X)

tb.Label(card, text="Label Text").pack(anchor=W)
label_entry = tb.Entry(card, font=("Segoe UI", 12))
label_entry.pack(fill=X, pady=5)
label_entry.bind("<Return>", lambda e: add_label())

tb.Label(card, text="Quantity (1–48)").pack(anchor=W, pady=(10, 0))
qty_entry = tb.Entry(card, font=("Segoe UI", 12))
qty_entry.pack(fill=X, pady=5)
qty_entry.bind("<Return>", lambda e: add_label())

add_btn = tb.Button(
    card,
    text="➕ Add Label",
    bootstyle=SUCCESS,
    command=add_label
)
add_btn.pack(fill=X, pady=12)

reset_btn = tb.Button(
    card,
    text="♻ Reset",
    bootstyle=DANGER,
    command=reset_app
)
reset_btn.pack(fill=X)

total_label = tb.Label(
    app,
    text="Total Labels: 0",
    font=("Segoe UI", 13, "bold"),
    bootstyle=WARNING
)
total_label.pack(pady=10)

run_btn = tb.Button(
    app,
    text="🚀 Replace & Open File",
    bootstyle=PRIMARY,
    width=30,
    command=run_replace
)
run_btn.pack(pady=15)

app.mainloop()
